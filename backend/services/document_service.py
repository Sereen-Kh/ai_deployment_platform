"""Document processing service for RAG."""

import os
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from uuid import UUID

import structlog
from fastapi import UploadFile
from sqlalchemy.ext.asyncio import AsyncSession

from core.config import settings
from api.models.rag import Document, DocumentStatus
from services.rag_service import RAGService

logger = structlog.get_logger()


class DocumentService:
    """Service for document processing and ingestion."""

    def __init__(self, db: AsyncSession):
        self.db = db
        self.rag_service = RAGService()

    async def process_document(
        self,
        file: UploadFile,
        collection_name: str,
        user_id: UUID,
        metadata: dict[str, Any] | None = None,
    ) -> Document:
        """Process and ingest a document."""
        # Read file content
        content = await file.read()
        file_size = len(content)

        # Validate size
        max_size = settings.upload_max_size_mb * 1024 * 1024
        if file_size > max_size:
            raise ValueError(
                f"File size exceeds maximum of {settings.upload_max_size_mb}MB")

        # Determine file type
        file_type = self._get_file_type(file.filename or "", file.content_type)

        # Create document record
        document = Document(
            name=file.filename or "unknown",
            file_path="",  # Will update after saving
            file_type=file_type,
            file_size=file_size,
            user_id=user_id,
            collection_name=collection_name,
            status=DocumentStatus.PROCESSING.value,
            metadata=metadata or {},
        )
        self.db.add(document)
        await self.db.commit()
        await self.db.refresh(document)

        try:
            # Save file temporarily for processing
            with tempfile.NamedTemporaryFile(
                delete=False, suffix=f".{file_type}"
            ) as tmp_file:
                tmp_file.write(content)
                tmp_path = tmp_file.name

            # Extract text based on file type
            text = await self._extract_text(tmp_path, file_type)

            # Chunk the text
            chunks = self._chunk_text(
                text=text,
                document_id=document.id,
                document_name=document.name,
                chunk_size=1000,
                chunk_overlap=200,
            )

            # Add chunks to vector store
            chunk_count = await self.rag_service.add_chunks(
                collection_name=collection_name,
                chunks=chunks,
            )

            # Update document record
            document.status = DocumentStatus.COMPLETED.value
            document.chunk_count = chunk_count
            document.token_count = len(text.split())  # Rough estimate
            document.processed_at = datetime.now(timezone.utc)

            # Clean up temp file
            os.unlink(tmp_path)

        except Exception as e:
            logger.error(
                "Document processing failed",
                document_id=str(document.id),
                error=str(e),
            )
            document.status = DocumentStatus.FAILED.value
            document.error_message = str(e)

        await self.db.commit()
        await self.db.refresh(document)

        return document

    def _get_file_type(self, filename: str, content_type: str | None) -> str:
        """Determine file type from filename or content type."""
        ext = Path(filename).suffix.lower().lstrip(".")

        type_map = {
            "pdf": "pdf",
            "docx": "docx",
            "doc": "docx",
            "txt": "txt",
            "md": "md",
            "markdown": "md",
            "csv": "csv",
            "xlsx": "xlsx",
            "xls": "xlsx",
            "json": "json",
            "html": "html",
            "htm": "html",
        }

        return type_map.get(ext, "txt")

    async def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text content from document."""
        if file_type == "pdf":
            return await self._extract_pdf(file_path)
        elif file_type == "docx":
            return await self._extract_docx(file_path)
        elif file_type in ["txt", "md"]:
            return await self._extract_text_file(file_path)
        elif file_type == "csv":
            return await self._extract_csv(file_path)
        elif file_type == "json":
            return await self._extract_json(file_path)
        else:
            # Fallback to text extraction
            return await self._extract_text_file(file_path)

    async def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error("PDF extraction failed", error=str(e))
            raise ValueError(f"Failed to extract text from PDF: {e}")

    async def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error("DOCX extraction failed", error=str(e))
            raise ValueError(f"Failed to extract text from DOCX: {e}")

    async def _extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    async def _extract_csv(self, file_path: str) -> str:
        """Extract text from CSV."""
        import csv

        rows = []
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(" | ".join(row))

        return "\n".join(rows)

    async def _extract_json(self, file_path: str) -> str:
        """Extract text from JSON."""
        import json

        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        return json.dumps(data, indent=2)

    def _chunk_text(
        self,
        text: str,
        document_id: UUID,
        document_name: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> list[dict[str, Any]]:
        """Split text into overlapping chunks."""
        chunks = []
        words = text.split()

        if len(words) <= chunk_size:
            # Single chunk for small documents
            chunks.append({
                "content": text.strip(),
                "document_id": str(document_id),
                "document_name": document_name,
                "chunk_index": 0,
                "metadata": {"total_chunks": 1},
            })
            return chunks

        start = 0
        chunk_index = 0

        while start < len(words):
            end = start + chunk_size
            chunk_words = words[start:end]
            chunk_text = " ".join(chunk_words)

            chunks.append({
                "content": chunk_text,
                "document_id": str(document_id),
                "document_name": document_name,
                "chunk_index": chunk_index,
                "metadata": {},
            })

            # Move start with overlap
            start = end - chunk_overlap
            chunk_index += 1

        # Update total_chunks in metadata
        for chunk in chunks:
            chunk["metadata"]["total_chunks"] = len(chunks)

        return chunks
